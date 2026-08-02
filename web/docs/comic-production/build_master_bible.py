from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parents[2]
OUTPUT = ROOT / "WKE_COMIC_MASTER_PRODUCTION_BIBLE.docx"
COVER_IMAGE = WORKSPACE / "work" / "comic-review" / "chapter-1" / "01-cover.png"
CHAPTER_2_PAGE_1 = ROOT / "source-pages" / "chapter-02" / "page-01.png"


# compact_reference_guide token map
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_DISTANCE = Inches(0.492)
FOOTER_DISTANCE = Inches(0.492)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

NAVY = RGBColor(22, 58, 95)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 99, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF6DC"
LIGHT_PURPLE = "F1EAF8"
LIGHT_GREEN = "E8F4E8"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_MARGIN_TOP_BOTTOM, start=CELL_MARGIN_START_END,
                     bottom=CELL_MARGIN_TOP_BOTTOM, end=CELL_MARGIN_START_END):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

    set_table_borders(table)


def shade_row(row, fill=LIGHT_BLUE):
    for cell in row.cells:
        set_cell_shading(cell, fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table_text(table, header=True, size=9.2):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, size=size, color=BLACK, bold=header and row_index == 0)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def add_numbering_definition(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullet(doc, text, bullet_num_id):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, bullet_num_id)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_step(doc, label, body, decimal_num_id):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, decimal_num_id)
    label_run = paragraph.add_run(label + " ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, size=11, color=BLACK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_label_paragraph(doc, label, body, fill=None):
    if fill:
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [CONTENT_WIDTH_DXA])
        set_cell_shading(table.cell(0, 0), fill)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
    else:
        paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(label + ": ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, size=11, color=BLACK)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    return paragraph


def add_picture_with_alt(doc, path, width, alt_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = MARGIN
        section.right_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.header_distance = HEADER_DISTANCE
        section.footer_distance = FOOTER_DISTANCE

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        hr = hp.add_run("WKE EDUCATIONAL COMIC SERIES  |  PRODUCTION BIBLE")
        set_run_font(hr, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        fr = fp.add_run("Version 0.1  |  Page ")
        set_run_font(fr, size=8.5, color=MUTED)
        add_page_field(fp)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet")

    # Editorial-cover pattern with a compact visual anchor.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("PRODUCTION REFERENCE  |  VERSION 0.1")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("WKE Educational Comic Series")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Master Production Bible")
    set_run_font(r, size=15, color=DARK_BLUE, italic=True)

    if COVER_IMAGE.exists():
        add_picture_with_alt(
            doc,
            COVER_IMAGE,
            2.65,
            "Chapter 1 cover showing Mia, Zara, Ethan, and Leo in the valley setting.",
        )
        add_caption(doc, "Chapter 1 visual anchor: A New Friend Arrives")

    add_label_paragraph(
        doc,
        "Purpose",
        "Protect story canon, educational quality, character identity, and visual continuity from script through lesson-player publication.",
        LIGHT_GOLD,
    )

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("Learners", "Ages 9-12 | CEFR A1-A2"),
        ("Confirmed cast", "Mia, Zara, Leo, Ethan, and Keelan"),
        ("Current status", "Chapter 1 published | Chapter 2 in production"),
        ("Updated", "August 2, 2026"),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
        table.cell(idx, 0).paragraphs[0].runs[0].bold = True
    set_table_geometry(table, [2700, 6660])
    style_table_text(table, header=False, size=9.5)

    add_page_break(doc)

    add_heading(doc, "1. How to use this bible", 1)
    add_label_paragraph(
        doc,
        "Production rule",
        "Read the relevant canon, character, language, and scene rules before producing a page. Resolve conflicts before generating later artwork.",
        LIGHT_BLUE,
    )
    usage_num_id = add_numbering_definition(doc, "decimal")
    add_step(doc, "Confirm outcomes.", "Define the story movement, character movement, and learner-language goal.", usage_num_id)
    add_step(doc, "Script the page.", "Record exact panel action, dialogue, positions, props, and transition from the prior page.", usage_num_id)
    add_step(doc, "Check canon.", "Compare names, identity colors, models, ship, crash-site map, and object locations.", usage_num_id)
    add_step(doc, "Create and letter.", "Generate or illustrate from approved references, then add only approved text.", usage_num_id)
    add_step(doc, "Review and publish.", "Complete learning and continuity QA, preserve alternates, and upload the single canonical page.", usage_num_id)

    add_heading(doc, "Canon hierarchy", 2)
    canon_num_id = add_numbering_definition(doc, "decimal")
    add_step(doc, "Owner decisions", "Newest explicit decisions confirmed by the project owner.", canon_num_id)
    add_step(doc, "Approved artwork", "Published dialogue and final canonical pages.", canon_num_id)
    add_step(doc, "Approved scripts", "The newest chapter and page records.", canon_num_id)
    add_step(doc, "Working standards", "Recommendations in this bible that have not yet been explicitly locked.", canon_num_id)
    add_step(doc, "Drafts", "Unapproved generations and alternates never overrule canon.", canon_num_id)

    add_heading(doc, "Series promise", 1)
    p = doc.add_paragraph(
        "Four friends discover Keelan, a stranded blue alien, and are drawn into a larger mystery. Their adventures combine curiosity, courage, friendship, science, discovery, and age-appropriate wonder."
    )
    add_bullet(doc, "The story must be engaging even when no lesson activity is attached.", bullet_num_id)
    add_bullet(doc, "Images should let learners infer meaning without translating every word.", bullet_num_id)
    add_bullet(doc, "Characters model cooperation, empathy, reasoning, and safe decision-making.", bullet_num_id)
    add_bullet(doc, "Every chapter gives teachers opportunities for prediction, retelling, speaking, and comprehension.", bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "2. Cast canon", 1)

    cast_table = doc.add_table(rows=6, cols=3)
    headers = ["Character", "Identity", "Model and story anchors"]
    for col, text in enumerate(headers):
        cast_table.cell(0, col).text = text
    cast_rows = [
        ("Mia", "Purple", "Long straight black hair; purple clip and dress; science-oriented, cautious, and analytical."),
        ("Zara", "Yellow", "Long dark curly hair; yellow headband and sun shirt; warm, encouraging, and empathetic."),
        ("Leo", "Blue", "Short black hair; blue wave shirt; energetic, playful, curious, and action-oriented."),
        ("Ethan", "Green", "Curly brown hair; green mountain shirt; imaginative, hopeful, adventurous, and persistent."),
        ("Keelan", "Blue alien", "Blue fur, pointed ears, golden-dark eyes, orange stars, fangs, tail; gentle, magical, and weak after the crash."),
    ]
    for row_idx, row_data in enumerate(cast_rows, start=1):
        for col_idx, value in enumerate(row_data):
            cast_table.cell(row_idx, col_idx).text = value
    shade_row(cast_table.rows[0])
    set_repeat_table_header(cast_table.rows[0])
    set_table_geometry(cast_table, [1600, 1700, 6060])
    style_table_text(cast_table, header=True, size=9.2)

    add_heading(doc, "Character voice", 2)
    add_label_paragraph(doc, "Mia", "Evidence-focused and thoughtful. She may use slightly more precise language but must remain A1-A2 accessible.", LIGHT_PURPLE)
    add_label_paragraph(doc, "Zara", "Supportive, expressive, and attentive to feelings.", LIGHT_GOLD)
    add_label_paragraph(doc, "Leo", "Direct, enthusiastic, and ready to act.", LIGHT_BLUE)
    add_label_paragraph(doc, "Ethan", "Optimistic, imaginative, and persistent.", LIGHT_GREEN)
    add_label_paragraph(doc, "Keelan", "Gentle and initially cautious; short statements reflect limited language and weakness.", LIGHT_BLUE)

    add_heading(doc, "Keelan model lock", 2)
    p = doc.add_paragraph(
        "Before additional final pages are approved, create a reference sheet that fixes Keelan’s height, ear proportions, star pattern, finger count, tail, fang placement, and magic effect."
    )
    p.paragraph_format.keep_with_next = True
    add_bullet(doc, "Use the same golden-dark eye design and orange star motif in every panel.", bullet_num_id)
    add_bullet(doc, "Keep a stable child-to-Keelan scale in group shots.", bullet_num_id)
    add_bullet(doc, "Define one approved paw anatomy and use it from every angle.", bullet_num_id)
    add_bullet(doc, "All human characters must clearly appear 9-12, with practical clothing and age-appropriate poses.", bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "3. Learner and language framework", 1)
    add_label_paragraph(doc, "Confirmed learner", "Ages 9-12 at CEFR A1-A2, reading independently or with teacher support.", LIGHT_GOLD)

    add_heading(doc, "Language controls", 2)
    for text in [
        "Prefer high-frequency concrete words that are supported by the picture.",
        "Keep most balloons between 3 and 10 words and focused on one idea.",
        "Keep narration concise, normally 8-18 words, and use it only for necessary transitions.",
        "Introduce approximately 6-10 useful target words or phrases per chapter.",
        "Repeat important sentence frames naturally in later panels and chapters.",
        "Avoid unexplained idioms, sarcasm, dense exposition, and abstract vocabulary.",
        "Use names when a pronoun or speaker could be unclear.",
        "Make every balloon tail unambiguous for emerging readers.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Language already established", 2)
    language_table = doc.add_table(rows=7, cols=2)
    language_table.cell(0, 0).text = "Function"
    language_table.cell(0, 1).text = "Published example"
    examples = [
        ("Reassuring", "I’m not going to hurt you."),
        ("Ability", "I can use a little magic."),
        ("Suggesting", "Maybe we should help him."),
        ("Imperatives", "Wait! / Run! / Stay together."),
        ("Possibility", "Maybe it powers the ship."),
        ("Past cause and present result", "My ship crashed. I’m weak."),
    ]
    for idx, row_data in enumerate(examples, start=1):
        language_table.cell(idx, 0).text = row_data[0]
        language_table.cell(idx, 1).text = row_data[1]
    shade_row(language_table.rows[0])
    set_repeat_table_header(language_table.rows[0])
    set_table_geometry(language_table, [3000, 6360])
    style_table_text(language_table, header=True, size=9.4)

    add_heading(doc, "Educational quality gate", 2)
    add_bullet(doc, "An A1 learner can follow the central action from images and familiar words.", bullet_num_id)
    add_bullet(doc, "An A2 learner can read most dialogue independently and explain key choices.", bullet_num_id)
    add_bullet(doc, "New language is inferable from action, expression, object, or repetition.", bullet_num_id)
    add_bullet(doc, "The page supports a simple retell, prediction, or speaking prompt.", bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "4. Visual and continuity standard", 1)
    add_label_paragraph(
        doc,
        "Working canvas",
        "1122 x 1402 pixels, portrait, approximately 4:5. Chapter 1 pages 3-6 currently differ and must be standardized before unified print or reader export.",
        LIGHT_BLUE,
    )

    add_heading(doc, "Art direction", 2)
    for text in [
        "Warm, detailed, painterly comic rendering with manga/anime influence.",
        "Saturated natural color, golden forest light, expressive faces, and readable body language.",
        "Three to six panels with consistent black borders and gutters.",
        "Parchment-style panel numbers, white balloons, black outlines, and clear tails.",
        "Yellow, black-outlined display lettering reserved for major sound effects.",
        "Stable identity system: Mia/purple, Zara/yellow-sun, Leo/blue-wave, Ethan/green-mountain, Keelan/blue-orange-stars.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Crash-site map", 2)
    add_label_paragraph(
        doc,
        "Spatial canon",
        "Chapter 1 places the broad saucer on the left and the root/tree hollow on the right. Each child keeps their own color-coded backpack; no duplicate loose packs remain beside the ship entrance.",
        LIGHT_GREEN,
    )
    add_label_paragraph(
        doc,
        "Current conflict",
        "Chapter 2 page 1 changes or blurs the ship/hollow relationship. Backpacks may be obscured by front-facing poses, but remain with their owners under scene-map canon v1.1.",
        LIGHT_GOLD,
    )

    add_heading(doc, "Spacecraft model", 2)
    for text in [
        "Broad, flattened circular saucer with a raised central dome.",
        "Dark segmented metal shell, blue illuminated details, and a round entrance.",
        "Interior contains a glowing blue crystal or power object.",
        "Damage and smoke may develop, but the silhouette and doorway remain recognizable.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Preferred lettering workflow", 2)
    p = doc.add_paragraph(
        "Generate clean artwork without final dialogue where practical, then add balloons, tails, panel numbers, and approved text in a controlled lettering pass. This prevents spelling drift and unclear speakers."
    )

    add_page_break(doc)
    add_heading(doc, "5. Chapter architecture", 1)
    architecture = [
        ("Hook or continuation.", "Reconnect readers to the unresolved question."),
        ("Immediate goal.", "Give the characters a clear problem, need, or decision."),
        ("Discovery.", "Reveal useful information about Keelan, the ship, the world, or a friend."),
        ("Choice or complication.", "Require cooperation, courage, empathy, or reasoning."),
        ("Small payoff.", "Resolve one question or complete one meaningful action."),
        ("Forward pull.", "End with a new question, discovery, risk, or promise."),
    ]
    architecture_num_id = add_numbering_definition(doc, "decimal")
    for label, body in architecture:
        add_step(doc, label, body, architecture_num_id)

    add_heading(doc, "Backward-design fields", 2)
    for text in [
        "Story goal and emotional/character goal.",
        "One communicative language goal.",
        "Six to ten target words or phrases.",
        "Grammar learners will encounter, without forcing explicit explanation into the comic.",
        "Visual supports for unfamiliar meaning.",
        "Three to five post-reading comprehension or speaking prompts.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Production workflow", 1)
    workflow = [
        ("Define outcomes.", "Decide what changes in the story and what learners will understand or use."),
        ("Write the beat sheet.", "Confirm hook, goal, discovery, complication, payoff, and forward pull."),
        ("Write page scripts.", "Record every panel’s composition, dialogue, emotion, props, and continuity anchors."),
        ("Review language.", "Check A1-A2 readability, dialogue length, visual support, and speaker clarity."),
        ("Review canon.", "Check models, identity colors, ship, location, props, and previous-page match."),
        ("Create art.", "Use approved prompt blocks, models, and scene maps."),
        ("Letter the page.", "Use only approved script text and verify reading order and tails."),
        ("Run QA.", "Complete the continuity and learning checklist at full resolution."),
        ("Archive and upload.", "Preserve alternates, mark one canonical file, update records, and publish in order."),
    ]
    workflow_num_id = add_numbering_definition(doc, "decimal")
    for label, body in workflow:
        add_step(doc, label, body, workflow_num_id)

    add_page_break(doc)
    add_heading(doc, "6. Current story canon", 1)
    add_heading(doc, "Chapter 1 - A New Friend Arrives", 2)
    add_label_paragraph(doc, "Status", "Published. Cover plus six story pages.", LIGHT_GREEN)
    p = doc.add_paragraph(
        "The friends debate whether aliens exist, witness a crash, investigate the ship and glowing crystal, follow movement to a tree hollow, and meet Keelan. The chapter ends when Keelan says, “Wait!”"
    )
    beats = [
        "Meadow play and science reading establish the four friends.",
        "Mia doubts aliens exist; Ethan refuses to give up hope.",
        "A spacecraft crashes near the woods.",
        "The children enter the saucer and find a glowing blue crystal.",
        "Movement and golden eyes appear in a nearby tree hollow.",
        "Keelan emerges and stops the children from running.",
    ]
    for beat in beats:
        add_bullet(doc, beat, bullet_num_id)

    add_heading(doc, "Chapter 2 - title TBD", 2)
    add_label_paragraph(doc, "Status", "In production. One canonical page confirmed.", LIGHT_GOLD)
    if CHAPTER_2_PAGE_1.exists():
        add_picture_with_alt(
            doc,
            CHAPTER_2_PAGE_1,
            3.55,
            "Chapter 2 page 1 showing the children meeting Keelan at the crash site.",
        )

    add_page_break(doc)
    add_heading(doc, "Chapter 2, Page 1 record", 2)
    page1 = [
        ("Panel 1", "Keelan calls, “Wait! Please don’t run!”"),
        ("Panel 2", "The children react: “It can talk…” and “That’s impossible.”"),
        ("Panel 3", "Keelan reassures them: “I’m not going to hurt you.”"),
        ("Panel 4", "Keelan introduces himself. A child observes that he is cute."),
        ("Panel 5", "Keelan reveals limited magic. A child responds, “Magic?!”"),
        ("Panel 6", "The children notice he is hurt; Keelan says the crash left him weak; they consider helping."),
    ]
    for label, body in page1:
        add_label_paragraph(doc, label, body)

    add_heading(doc, "Canon established", 2)
    for text in [
        "The alien’s name is Keelan and the approved dialogue uses he/him pronouns.",
        "Keelan is peaceful, speaks the children’s language, and can use limited magic.",
        "The crash left Keelan weak.",
        "The children begin moving from fear toward empathy and cooperation.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Next-page requirements", 2)
    for text in [
        "State the specific help Keelan needs.",
        "Preserve all five models and the approved crash-site map.",
        "Continue one clear A1-A2 communicative goal.",
        "Move toward a concrete Chapter 2 payoff and ending hook.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "7. Approval checklist", 1)

    checklist_sections = [
        ("Story connection", [
            "First panel matches prior positions, emotion, props, and lighting.",
            "The page has one clear story job and understandable transitions.",
            "The last panel creates a reason to continue.",
        ]),
        ("Character and environment", [
            "All models, identity colors, clothing symbols, age, and anatomy match.",
            "Ship, hollow, landmarks, backpacks, books, and crystal remain consistent.",
            "Screen direction and character order do not jump without explanation.",
        ]),
        ("Language and learning", [
            "A1 learners can follow the main action from the visual context.",
            "A2 learners can read most dialogue independently.",
            "Every balloon has one clear idea, an obvious speaker, and approved text.",
        ]),
        ("Technical", [
            "Canvas, gutters, panel numbers, lettering size, reading order, and safe area pass.",
            "No unintended extra fingers, duplicate characters, cropped limbs, or floating objects.",
            "Canonical filename, page order, alternate archive, and chapter record are correct.",
        ]),
    ]
    for heading, items in checklist_sections:
        add_heading(doc, heading, 2)
        for item in items:
            add_bullet(doc, "[ ] " + item, bullet_num_id)

    add_heading(doc, "Known issues requiring resolution", 1)
    issues_table = doc.add_table(rows=9, cols=3)
    for col, header in enumerate(("Priority", "Issue", "Required decision")):
        issues_table.cell(0, col).text = header
    issues = [
        ("High", "Mixed page aspect ratios", "Approve one canvas and standardize all pages."),
        ("High", "Ship design and crash-site geography drift", "Approve a spacecraft model and scene map."),
        ("Medium", "Backpacks become obscured", "Keep them with their owners and do not add loose duplicates."),
        ("Medium", "Keelan model drifts", "Approve the full model sheet before more final pages."),
        ("Medium", "Magic appears before its dialogue reveal", "Move the effect or revise the script beat."),
        ("Low", "Mia hair-clip count varies", "Lock one model."),
        ("Low", "Chapter 1 page 5 lacks panel numbers", "Add numbers or revise the series convention."),
        ("Low", "Some balloon tails are unclear", "Clarify attribution during lettering."),
    ]
    for row_idx, row_data in enumerate(issues, start=1):
        for col_idx, value in enumerate(row_data):
            issues_table.cell(row_idx, col_idx).text = value
    shade_row(issues_table.rows[0])
    set_repeat_table_header(issues_table.rows[0])
    set_table_geometry(issues_table, [1200, 3000, 5160])
    style_table_text(issues_table, header=True, size=8.8)

    add_heading(doc, "Open story decisions", 2)
    for text in [
        "Chapter 2 title.",
        "The specific repair, object, or action Keelan needs.",
        "The Chapter 2 payoff and final hook.",
        "Whether the children help alone, retrieve the crystal, or seek an adult.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "Change log", 2)
    add_label_paragraph(
        doc,
        "Version 0.1 - August 2, 2026",
        "Confirmed the cast names and color identities, Keelan’s name, the ages 9-12 learner band, CEFR A1-A2 target, Chapter 1 record, Chapter 2 page 1, and the first production standards.",
        LIGHT_BLUE,
    )

    # Core properties and final geometry pass.
    doc.core_properties.title = "WKE Educational Comic Series - Master Production Bible"
    doc.core_properties.subject = "Canon, learning design, visual continuity, and production workflow"
    doc.core_properties.author = "We Know English Center"
    doc.core_properties.keywords = "comic, education, A1, A2, continuity, production bible"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
